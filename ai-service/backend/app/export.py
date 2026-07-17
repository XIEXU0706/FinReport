"""
    Word 文档导出工具
    将 AI 生成的 markdown 报告转换为 .docx 文件
"""
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def markdown_to_docx(title: str, md: str) -> BytesIO:
    """将 markdown 文本转换为 .docx 文件，返回 BytesIO 对象"""
    doc = Document()

    # 设置默认中文字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 标题
    para = doc.add_heading(title, level=0)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 逐行解析 markdown
    lines = md.split('\n')
    i = 0
    in_table = False
    table_data = []
    table_header_skipped = False

    while i < len(lines):
        line = lines[i]

        # 表格行
        if line.startswith('|') and line.endswith('|'):
            parts = [p.strip() for p in line.split('|')[1:-1]]
            if not in_table:
                in_table = True
                table_data = [parts]
                table_header_skipped = False
                i += 1
                continue
            # 分隔行 |---|---| 跳过
            if re.match(r'^[\s\|:\-]+$', line) and not table_header_skipped:
                table_header_skipped = True
                i += 1
                continue
            table_data.append(parts)
            i += 1
            continue
        else:
            # 表格结束，写入文档
            if in_table and table_data:
                _add_table(doc, table_data)
                in_table = False
                table_data = []
                table_header_skipped = False

        # 空行 → 段落间距
        if not line.strip():
            doc.add_paragraph('')
            i += 1
            continue

        # 标题 ## 或 ###
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1)) - 1  # ## → level 1, ### → level 2
            text = heading_match.group(2)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # 列表项 - 或 *
        list_match = re.match(r'^[\-\*]\s+(.+)$', line)
        if list_match:
            text = _inline_md_to_runs(list_match.group(1))
            p = doc.add_paragraph(style='List Bullet')
            _add_runs(p, text)
            i += 1
            continue

        # 数字列表 1. 2.
        num_match = re.match(r'^\d+\.\s+(.+)$', line)
        if num_match:
            text = _inline_md_to_runs(num_match.group(1))
            p = doc.add_paragraph(style='List Number')
            _add_runs(p, text)
            i += 1
            continue

        # 普通段落
        text = _inline_md_to_runs(line)
        p = doc.add_paragraph()
        _add_runs(p, text)
        i += 1

    # 处理末尾未关闭的表格
    if in_table and table_data:
        _add_table(doc, table_data)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _inline_md_to_runs(text: str):
    """解析行内 markdown 标记，返回 (text, bold) 元组列表"""
    parts = []
    # 匹配 **bold** 和 `code`
    pattern = r'(\*\*(.+?)\*\*|`(.+?)`)'
    last = 0
    for m in re.finditer(pattern, text):
        if m.start() > last:
            parts.append((text[last:m.start()], False, False))
        if m.group(2):   # **bold**
            parts.append((m.group(2), True, False))
        elif m.group(3): # `code`
            parts.append((m.group(3), False, True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False, False))
    return parts if parts else [(text, False, False)]


def _add_runs(paragraph, parts):
    """将 (text, bold, code) 元组列表添加到段落"""
    for text, bold, is_code in parts:
        run = paragraph.add_run(text)
        run.bold = bold
        if is_code:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        else:
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(11)


def _add_table(doc, data):
    """添加表格到文档"""
    if not data:
        return
    rows = len(data)
    cols = max(len(r) for r in data) if data else 0
    if cols == 0:
        return
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'
    for r_idx, row_data in enumerate(data):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = row_data[c_idx] if c_idx < len(row_data) else ''
            # 表头加粗
            if r_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
